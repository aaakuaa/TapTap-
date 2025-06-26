#文件
import os
#爬虫
import requests
from bs4 import BeautifulSoup
#word
from docx import Document
from docx.shared import Inches
#excel
from openpyxl import Workbook
from openpyxl.styles import Alignment
from openpyxl.drawing.image import Image as XLImage
#时间
import time
#多线程
import threading
#计数
from collections import Counter
#可视化——绘图框架
import matplotlib.pyplot as plt

# 创建文件夹
os.makedirs("D:\\爬", exist_ok=True)  # 创建一个名为“爬”的文件夹
os.makedirs("D:\\爬\\图", exist_ok=True)  # 创建一个名为“图”的文件夹，用于保存游戏图片

# 创建Word文档
doc = Document()  # 创建一个新的Word文档对象
doc.add_heading('TapTap热玩榜', level=1)  # 添加标题到文档

# 创建Excel文件并设置格式
wb = Workbook()  # 创建一个新的Excel工作簿对象
ws = wb.active  # 获取活动工作表
ws.append(["排名", "名称", "", "评分","类型"])  # 添加Excel表头
ws.column_dimensions['B'].width = 20  # 设置列宽
ws.column_dimensions['C'].width = 12

All_tags = []  # 用于存储所有游戏标签的列表
lock = threading.Lock()  # 创建一个线程锁，用于在多线程中同步访问共享资源
results = []  # 用于存储所有结果的列表

# 定义一个函数，用于从指定页面抓取数据
def fetch_page_data(page):
    header = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36 Edg/124.0.0.0"
    }  # 反爬虫——设置请求头，模拟浏览器访问

    response = requests.get(f"https://www.taptap.cn/top/played?page={page}", headers=header)  # 发送GET请求获取页面内容
    if response.ok:  # 如果请求成功
        print(f"请求成功 - Page {page}")
        html = response.text
        soup = BeautifulSoup(html, "html.parser")  # 使用BeautifulSoup解析页面内容

        links = soup.findAll("a")  # 查找所有<a>标签
        printed_links = []  # printed_links——寄存目标网页所搜寻的每款游戏的超链接，用于进行进一步跟踪爬取游戏tag
        for link in links:  # 储存过程
            href = link.get('href')
            if href.startswith('/app') and href not in printed_links and len(printed_links) < 10:
                printed_links.append(href)

        all_titles = soup.findAll("span", attrs={"class": "text text-default--size"})  # 查找所有游戏标题
        all_images = soup.findAll("img", attrs={"class": "tap-image app-icon__img"})  # 查找所有游戏图片
        all_ratings = soup.findAll("div", attrs={"class": "tap-rating__number rate-number-font"})  # 查找所有游戏评分

        for idx, (title, image, rating) in enumerate(zip(all_titles, all_images, all_ratings), start=1):#  打包（zip）内容，枚举（enumerate）迭代赋值
            name = title.string.strip()     #  提取文本
            image_url = image["src"]    #提取图片地址，以便后续下载
            rating_value = rating.string.strip() if rating else "暂无评分"  # 防止部分没有评分，特殊说明

            print(f"找到游戏: {name}")
            image_response = requests.get(image_url)    # 追踪图片地址
            if image_response.ok:
                image_file_path = os.path.join("D:\爬\图", f"{name}.jpg")  # 构建图片保存路径
                with open(image_file_path, "wb") as f:
                    f.write(image_response.content)  # 将图片内容写入文件
                print(f"保存信息: {name}")

                ranking = (page - 1) * 10 + idx     #  根据网页页数计算排名

                response = requests.get(f"https://www.taptap.cn{printed_links[idx-1]}", headers=header)     #  爬取已经存储好的链接，进行追踪爬取游戏tag
                html1 = response.text
                soup = BeautifulSoup(html1, 'html.parser')
                tags = soup.find_all('a', class_='tap-router tap-chip tap-chip--leading tap-chip--default')     # 找到保存tag所在信息

                with lock:          #创建锁lock，确保全局共享资源被正确访问使用
                    game_tags = []      #创建空列表来存储tag
                    for tag in tags:
                        text = tag.get_text(strip=True)     #提取tag（strip=True——消除空白，提取文本）
                        All_tags.append(text)       #tag大合集，用以后续统计tag做可视化
                        game_tags.append(text)      #单独游戏本体tag，用于制作文件
                    # 将爬取的信息打包，以便后续使用, ', '.join(game_tags)——用逗号作间隔将游戏本体tag缝成字符串
                    results.append((ranking, name, rating_value, image_file_path, ', '.join(game_tags)))
            else:
                print(f"下载失败 '{name}'")
    else:
        print(f"请求失败 - Page {page}")

    time.sleep(1)  # 休眠1秒，避免对服务器造成过大负载

# 创建并启动多线程，每个线程负责抓取一个页面的数据
threads = []    #创建列表用于存放多个线程程序
for page in range(1, 16):
    #创建多线程程序，执行函数（target）为上文fetch_page_data，参数（args，元组）为page
    thread = threading.Thread(target=fetch_page_data, args=(page,))
    threads.append(thread)       #保存进程
    thread.start()               #开始进程

# 等待所有线程完成
for thread in threads:
    thread.join()   #阻止后续程序进行，用join（）等待爬取进程结束

# 按排名排序结果，整理数据
results.sort(key=lambda x: x[0])

# 将结果添加到Word文档和Excel文件中
for result in results:
    ranking, name, rating_value, image_file_path, game_tags = result    #导入数据

    #依次将数据编入word文档
    doc.add_paragraph(f"排名: {ranking}")
    doc.add_paragraph(f"名称: {name}")
    doc.add_paragraph(f"评分: {rating_value}")
    doc.add_paragraph(f"类型: {game_tags}")
    #设置图片规格，插入word
    doc.add_picture(image_file_path, width=Inches(1.0))
    doc.add_paragraph()
    doc.add_paragraph()

    #依次将数据写入excel文档
    img = XLImage(image_file_path)                  #用XLImage处理图片
    ws.add_image(img, f"C{ranking + 1}")            #在ranking+1行的C列插入图片
    ws.row_dimensions[ranking + 1].height = 68      #设置数据行的高度确保图片能装下
    #导入文本数据
    ws.cell(row=ranking + 1, column=1).value = ranking
    ws.cell(row=ranking + 1, column=2).value = name
    ws.cell(row=ranking + 1, column=4).value = rating_value
    ws.cell(row=ranking + 1, column=5).value = game_tags

# 设置Excel单元格对齐方式（居中对齐）
for row in ws.iter_rows():
    for cell in row:
        cell.alignment = Alignment(vertical='center')

# 保存Word文档和Excel文件
doc.save('D:\\爬\\TapTap热玩榜.docx')
wb.save('D:\\爬\\TapTap热玩榜.xlsx')

# 统计游戏标签出现次数并绘制柱状图
tag_counts = Counter(All_tags)      #用Counter方法统计tag和对应出现次数

plt.rcParams['font.sans-serif'] = ['SimHei']  # 设置中文字体为黑体，以防乱码

C_tags = list(tag_counts.keys())     #提取tag类型
C_counts = list(tag_counts.values()) #提取tag次数

tags, counts = zip(*sorted(zip(C_tags, C_counts), key=lambda x: x[1], reverse=True))    #按照次数高低排序并且赋值

tags = tags[:15]    #提取出现次数前15的tag
counts = counts[:15]

plt.figure(figsize=(12, 6))     #设置视图窗口大小
bars = plt.bar(tags, counts)    #tags作为x轴数据，counts作为y轴数据

for bar, count in zip(bars, counts):
    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height(), count, ha='center', va='bottom')
    # 在每个条形图的顶部居中显示其高度值。'bar.get_x（） + bar.get_width（） / 2'
    # 计算了条形图的中心位置作为文本的横坐标，'bar.get_height（）'
    # 获取了条形图的高度作为文本的纵坐标，'count'（次数）
    # 是要显示的文本内容
    # 'ha=center'表示水平居中对齐
    # 'va=bottom'表示垂直底部对齐


plt.yticks(range(0, max(counts) + 5, 5))       #设置y轴刻度为5，峰值为最大次数加5
plt.title('最受喜爱tags top15')                  #设置标题
plt.xlabel('游戏标签')                           #设置x轴属性
plt.ylabel('出现次数')                           #设置y轴属性

plt.tight_layout()      #生成柱状图
plt.show()              #展示柱状图
